# ==============================================================================
# report_generator.py
#
# This module provides the ReportGenerator class, which is responsible for
# creating shareable documents from the results of an analysis session. By
# separating this functionality, we keep the core analysis and visualisation
# code clean and focused on their respective tasks.
#
# Author: Tim Vos
# Last Modified: 2 October 2025
# ==============================================================================

import docx
from docx.shared import Inches
import io
from typing import List
from datetime import datetime

class ReportGenerator:
    """
    Generates reports from analysis artifacts, such as cached plot images.
    
    This class is a stateless service; it does not hold any data itself. It simply
    provides the functionality to process a given set of inputs (the plot cache)
    and produce a file as output.
    """

    def create_word_document(
        self,
        plot_cache: List[io.BytesIO],
        filename: str | None = None,
        analyzer_name: str | None = None,
        df_preview: pd.DataFrame | None = None
    ) -> None:
        """
        Creates a Microsoft Word report from cached plots and optional DataFrame preview.
    
        Args:
            plot_cache (List[io.BytesIO]): Cached plot images.
            filename (str | None): Optional target filename (.docx). If None, one will
                                   be automatically generated using analyzer_name + timestamp.
            analyzer_name (str | None): Logical name of the analyzer instance.
            df_preview (pd.DataFrame | None): Optional DataFrame preview to include.
        """
        # --- Handle filename creation ---
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M")
        base_name = analyzer_name or "report"
        if filename is None:
            filename = f"{base_name}_{timestamp}.docx"
        elif not filename.endswith(".docx"):
            filename += ".docx"
    
        # --- Start document ---
        document = docx.Document()
        document.add_heading(f"{base_name.title()} Report", level=1)
        document.add_paragraph(f"Generated on: {datetime.now():%A, %d %B %Y, %I:%M %p}")
    
        # --- Optional preview of df.head() ---
        if df_preview is not None:
            document.add_paragraph("Data preview (first 5 rows):")
            head = df_preview.head()
            table = document.add_table(rows=1, cols=len(head.columns))
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(head.columns):
                hdr_cells[i].text = str(col)
            for _, row in head.iterrows():
                row_cells = table.add_row().cells
                for i, val in enumerate(row):
                    row_cells[i].text = str(val)
            document.add_paragraph()  # space after table
    
        # --- Insert plots ---
        for i, img in enumerate(plot_cache):
            document.add_heading(f"Plot {i + 1}", level=2)
            document.add_picture(img, width=Inches(6.0))
            document.add_page_break()
    
        # --- Save ---
        try:
            document.save(filename)
            print(f"Successfully generated report: '{filename}'")
        except Exception as e:
            print(f"Error: Could not save the document. Reason: {e}")
